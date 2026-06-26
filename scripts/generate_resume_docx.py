from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


OUTPUT_PATH = "outputs/端侧AI模型部署方向_简历_v1.docx"


def set_run_font(run, font_name="Microsoft YaHei", font_size=9):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, font_size=11)
    run.bold = True
    return p


def add_normal_line(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0

    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1)
        r1.bold = True

        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0

    r = p.add_run("• " + text)
    set_run_font(r)
    return p


def main():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.52)
    section.right_margin = Inches(0.52)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(9)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("XXX")
    set_run_font(r, font_size=15)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(
        "电话：XXX ｜ 邮箱：XXX ｜ GitHub：https://github.com/mengpipi1019-create/edge-ai-onnxruntime-benchmark"
    )
    set_run_font(r, font_size=8.5)

    # Education
    add_heading(doc, "教育背景")
    add_normal_line(doc, "上海大学｜硕士｜电子信息工程相关方向｜20XX.09 - 至今")
    add_bullet(doc, "研究方向：低轨卫星边缘智能、联邦学习、强化学习资源优化；关注端侧 AI、边缘计算、模型部署、通信 AI。")
    add_normal_line(doc, "上海大学｜本科｜电子信息工程相关方向｜20XX.09 - 20XX.06")

    # Skills
    add_heading(doc, "专业技能")
    add_bullet(doc, "编程语言：熟悉 Python，能够完成模型推理、实验脚本、数据处理和结果可视化。")
    add_bullet(doc, "深度学习框架：熟悉 PyTorch 基础使用，掌握模型加载、前向推理、图像预处理和结果分析流程。")
    add_bullet(doc, "模型部署：了解 ONNX 与 ONNX Runtime，已完成 PyTorch 模型导出 ONNX、ONNX Runtime 推理部署、输出一致性验证和推理性能 benchmark。")
    add_bullet(doc, "数据处理与可视化：熟悉 NumPy、Pandas、Matplotlib，能够生成 CSV 实验结果和性能曲线图。")
    add_bullet(doc, "科研方向：了解联邦学习、异步聚合、强化学习资源优化、边缘智能和低轨卫星通信场景。")
    add_bullet(doc, "工具使用：熟悉 Git 基础操作，能够进行本地版本管理、GitHub 仓库上传和项目文档维护。")

    # Project 1
    add_heading(doc, "项目经历")
    add_normal_line(doc, "基于 PyTorch 与 ONNX Runtime 的轻量化图像分类模型部署与推理优化系统")
    add_normal_line(doc, "GitHub: https://github.com/mengpipi1019-create/edge-ai-onnxruntime-benchmark")
    add_normal_line(doc, "技术栈：Python、PyTorch、torchvision、ONNX、ONNX Runtime、NumPy、Pandas、Matplotlib、Pillow、Git")

    add_bullet(doc, "围绕端侧 AI 模型部署场景，构建 PyTorch → ONNX → ONNX Runtime 的图像分类推理流程，覆盖模型加载、真实图片预处理、ONNX 导出、部署推理、性能测试与可视化分析。")
    add_bullet(doc, "基于 ImageNet 预训练 ResNet18 完成真实图片 Top-K 分类推理，并统计 CPU 环境下的平均推理延迟、FPS 和 batch size 性能变化。")
    add_bullet(doc, "使用 torch.onnx.export 导出 ONNX 模型，并结合 onnx.checker、Top-K 分类一致性和 logits 最大绝对误差验证模型转换正确性，误差稳定在 1e-5 量级。")
    add_bullet(doc, "设计 batch size = 1、2、4、8 的推理 benchmark，生成 CSV 结果和性能曲线；实验显示 ONNX Runtime 在 CPU 上相比 PyTorch 实现约 2.00x~2.67x 推理加速。")

    # Project 2
    add_normal_line(doc, "面向低轨卫星边缘智能的异步联邦学习与资源优化系统")
    add_normal_line(doc, "项目类型：科研项目 / 论文项目")
    add_normal_line(doc, "技术栈：Python、PyTorch、联邦学习、强化学习、低轨卫星通信、资源优化")

    add_bullet(doc, "面向低轨卫星星座中星地链路间歇可见、通信资源受限和客户端更新异步到达等问题，研究适用于低轨卫星边缘智能场景的异步联邦学习架构。")
    add_bullet(doc, "构建包含本地训练、星间协同、星地上传和地面异步聚合的多层次训练流程，分析不同通信拓扑和资源控制策略对模型精度、时延和能耗的影响。")
    add_bullet(doc, "设计包含训练时延、通信能耗和模型精度的综合优化目标，并使用强化学习方法对通信功率、聚合频率和训练轮次等控制变量进行优化。")
    add_bullet(doc, "通过多组仿真实验对比不同架构和控制策略在非独立同分布数据、可见性受限和通信资源受限条件下的性能差异。")

    # Algorithm
    add_heading(doc, "算法与编程训练")
    add_bullet(doc, "已学习 LeetCode Hot 100 高频题，覆盖数组、链表、栈、队列、二叉树、动态规划、回溯、图论、堆和单调栈等模块。")
    add_bullet(doc, "能够使用 Python 实现常见数据结构与算法题解，并分析时间复杂度和空间复杂度。")

    # Positioning
    add_heading(doc, "求职定位")
    add_bullet(doc, "电子信息工程背景，研究方向聚焦低轨卫星边缘智能、联邦学习与通信资源优化，同时补充了 PyTorch、ONNX Runtime 和模型部署工程项目；希望应聘端侧 AI、模型部署、边缘智能、AI 工程化或通信 AI 相关实习岗位。")

    doc.save(OUTPUT_PATH)
    print(f"Saved resume to: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()